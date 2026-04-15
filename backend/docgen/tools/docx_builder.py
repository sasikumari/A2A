"""python-docx document assembler for cover pages, TOCs, styled content, tables, and images."""
from __future__ import annotations

import logging
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.oxml.ns import nsmap

from docgen.config import settings
from docgen.models import GeneratedContent, TableData

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Color constants (RGBColor)
# ---------------------------------------------------------------------------

BLUE_DARK = RGBColor(0x1E, 0x40, 0xAF)      # heading 1
NAVY = RGBColor(0x1E, 0x3A, 0x5F)            # heading 2
DARK_GRAY = RGBColor(0x37, 0x41, 0x51)       # heading 3 / body
BLUE_HEADER = RGBColor(0x1D, 0x4E, 0xD8)     # table header bg
LIGHT_BLUE = RGBColor(0xDB, 0xEA, 0xFE)      # alt row
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x0A, 0x0A, 0x0A)
BORDER_BLUE = RGBColor(0x3B, 0x82, 0xF6)


def _rgb_hex(color: RGBColor | str | Iterable[int]) -> str:
    """Normalize python-docx colors and raw RGB values to a six-char uppercase hex string."""
    if isinstance(color, str):
        return color.replace("#", "").upper()

    rgb_value = getattr(color, "rgb", None)
    if rgb_value is not None:
        return _rgb_hex(rgb_value)

    try:
        return bytes(color).hex().upper()
    except TypeError as exc:
        raise TypeError(f"Unsupported color value: {color!r}") from exc


# ---------------------------------------------------------------------------
# XML helpers
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_margins(cell, top: int = 60, bottom: int = 60, left: int = 108, right: int = 108):
    """Set table cell margins in twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcMar"))
    if existing is not None:
        tcPr.remove(existing)
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)


def _add_toc_field(doc: Document):
    """Insert a Word TOC field that updates when the user opens the document."""
    para = doc.add_paragraph(style="Normal")
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_paragraph_spacing(para, before=0, after=0)
    run = para.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar_separate = OxmlElement("w:fldChar")
    fldChar_separate.set(qn("w:fldCharType"), "separate")
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_separate)
    placeholder = OxmlElement("w:t")
    placeholder.text = "Update field to refresh entries and page numbers."
    run._r.append(placeholder)
    run._r.append(fldChar_end)


def _enable_update_fields_on_open(doc: Document):
    settings_el = doc.settings.element
    existing = settings_el.find(qn("w:updateFields"))
    if existing is not None:
        settings_el.remove(existing)
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings_el.append(update)


def _add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def _set_paragraph_spacing(para, before: int = 0, after: int = 6, line_rule=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line_rule:
        spacing.set(qn("w:lineRule"), line_rule)
    existing = pPr.find(qn("w:spacing"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(spacing)


def _add_header_border(para):
    """Add a blue bottom border to the header paragraph."""
    pPr = para._p.get_or_add_pPr()
    existing = pPr.find(qn("w:pBdr"))
    if existing is not None:
        pPr.remove(existing)
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), _rgb_hex(BLUE_DARK))
    pBdr.append(bottom)
    pPr.append(pBdr)


def _move_paragraph_after(doc: Document, para, after_idx: int):
    """
    Move a paragraph element to immediately after the paragraph at after_idx.
    Uses the live doc.paragraphs list (re-evaluated after each move).
    """
    body = doc.element.body
    paragraphs = doc.paragraphs
    if not paragraphs:
        return
    # Remove from current position
    body.remove(para._element)
    # Re-evaluate paragraph list (now one shorter)
    paragraphs = doc.paragraphs
    if after_idx < len(paragraphs):
        paragraphs[after_idx]._element.addnext(para._element)
    else:
        body.append(para._element)


def _split_long_paragraph(text: str, max_sentences: int = 4) -> list[str]:
    """Split a long paragraph into chunks of at most max_sentences sentences."""
    import re
    sentences = re.split(r'(?<=[.!?])\s+', text.strip())
    sentences = [s.strip() for s in sentences if s.strip()]
    if len(sentences) <= max_sentences:
        return [text]
    chunks = []
    for i in range(0, len(sentences), max_sentences):
        chunk = " ".join(sentences[i : i + max_sentences])
        if chunk:
            chunks.append(chunk)
    return chunks


# ---------------------------------------------------------------------------
# Document builder
# ---------------------------------------------------------------------------

class DocxBuilder:
    def __init__(self):
        self.doc = Document()
        self._setup_page()
        self._setup_default_style()
        _enable_update_fields_on_open(self.doc)

    def _setup_page(self):
        section = self.doc.sections[0]
        # US Letter
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        # 1-inch margins
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        # Suppress header/footer on cover page (first page)
        section.different_first_page_header_footer = True

    def _setup_default_style(self):
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = settings.default_font
        font.size = Pt(settings.default_font_size)
        font.color.rgb = BLACK
        self._setup_heading_styles()
        self._setup_toc_styles()

    def _setup_heading_styles(self):
        heading_specs = {
            "Heading 1": (18, BLUE_DARK, 18, 6),
            "Heading 2": (14, NAVY, 14, 6),
            "Heading 3": (12, DARK_GRAY, 12, 4),
        }
        for style_name, (size, color, before, after) in heading_specs.items():
            style = self.doc.styles[style_name]
            font = style.font
            font.name = settings.default_font
            font.size = Pt(size)
            font.bold = True
            font.color.rgb = color
            paragraph_format = style.paragraph_format
            paragraph_format.space_before = Pt(before)
            paragraph_format.space_after = Pt(after)
            paragraph_format.keep_with_next = True

    def _paragraph_style(self, name: str):
        try:
            return self.doc.styles[name]
        except KeyError:
            return self.doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    def _setup_toc_styles(self):
        try:
            toc_heading = self.doc.styles["TOC Heading"]
        except KeyError:
            toc_heading = self.doc.styles.add_style("TOC Heading", WD_STYLE_TYPE.PARAGRAPH)
        toc_heading.font.name = settings.default_font
        toc_heading.font.size = Pt(18)
        toc_heading.font.bold = True
        toc_heading.font.color.rgb = BLUE_DARK
        toc_hf = toc_heading.paragraph_format
        toc_hf.space_after = Pt(8)
        toc_hf.keep_with_next = True

        # python-docx default template omits TOC 1–3; Word uses them when the TOC field updates.
        entry_specs = [
            ("TOC 1", Pt(12), True, BLUE_DARK, Inches(0), Pt(10), Pt(4), Inches(6.25)),
            ("TOC 2", Pt(11), False, NAVY, Inches(0.22), Pt(6), Pt(2), Inches(6.03)),
            ("TOC 3", Pt(10), False, DARK_GRAY, Inches(0.44), Pt(4), Pt(2), Inches(5.81)),
        ]
        for name, size, bold, color, left_ind, before, after, tab_pos in entry_specs:
            st = self._paragraph_style(name)
            st.font.name = settings.default_font
            st.font.size = size
            st.font.bold = bold
            st.font.color.rgb = color
            pf = st.paragraph_format
            pf.left_indent = left_ind
            pf.space_before = before
            pf.space_after = after
            pf.tab_stops.clear_all()
            pf.tab_stops.add_tab_stop(tab_pos, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

    # ------------------------------------------------------------------
    # Cover page
    # ------------------------------------------------------------------

    # Revision history column specs per doc type
    _BRD_REV_HEADERS = ["Sr. No.", "Version No.", "Date of Change", "Change By", "Reviewed By", "Remarks"]
    _TSD_REV_HEADERS = ["Sr. No.", "Version", "Document Name", "Date of Change", "Remarks"]

    def _revision_row_data(self, doc_type: str, ri: int, revision: dict) -> list[str]:
        dt = doc_type.strip().lower()
        if dt == "brd":
            return [
                str(ri + 1),
                revision.get("version_no", revision.get("version", "")),
                revision.get("date_of_change", revision.get("date", "")),
                revision.get("changed_by", revision.get("change_by", revision.get("author", ""))),
                revision.get("reviewed_by", ""),
                revision.get("remarks", ""),
            ]
        else:  # TSD, Product Note, etc.
            return [
                str(ri + 1),
                revision.get("version_no", revision.get("version", "")),
                revision.get("document_name", ""),
                revision.get("date_of_change", revision.get("date", "")),
                revision.get("remarks", ""),
            ]

    def add_cover_page(self, title: str, subtitle: str = "", doc_type: str = "BRD",
                       version: str = "1.0", meta: dict | None = None):
        meta = meta or {}

        # ── Cover page (page 1) ───────────────────────────────────────────
        for _ in range(6):
            self.doc.add_paragraph()

        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(title_para, before=0, after=12)
        run = title_para.add_run(title)
        run.font.name = settings.default_font
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.color.rgb = BLUE_DARK

        if subtitle:
            sub_para = self.doc.add_paragraph()
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_paragraph_spacing(sub_para, before=0, after=8)
            run = sub_para.add_run(subtitle)
            run.font.name = settings.default_font
            run.font.size = Pt(18)
            run.font.bold = False
            run.font.color.rgb = NAVY

        type_para = self.doc.add_paragraph()
        type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(type_para, before=0, after=6)
        run = type_para.add_run(doc_type.upper())
        run.font.name = settings.default_font
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = DARK_GRAY

        ver_para = self.doc.add_paragraph()
        ver_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(ver_para, before=0, after=24)
        classification = meta.get("classification")
        version_line = f"Version {meta.get('version_number', version)}"
        if classification:
            version_line = f"{version_line}  |  {classification}"
        run = ver_para.add_run(version_line)
        run.font.name = settings.default_font
        run.font.size = Pt(12)
        run.font.color.rgb = DARK_GRAY

        # ── Page break → Revision History (page 2) ───────────────────────
        self.doc.add_page_break()

        rev_label = self.doc.add_paragraph()
        rev_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(rev_label, before=0, after=10)
        run = rev_label.add_run("Revision History")
        run.font.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = BLUE_DARK

        # Choose column headers based on doc type
        dt_norm = doc_type.strip().lower()
        if dt_norm == "brd":
            rev_headers = self._BRD_REV_HEADERS
        else:
            rev_headers = self._TSD_REV_HEADERS

        n_rev_cols = len(rev_headers)
        default_rev = [{"version": meta.get("version_number", version), "date_of_change": "", "remarks": "Initial draft"}]
        revision_rows = meta.get("revision_history") or default_rev
        required_rows = max(2, len(revision_rows))

        rev_table = self.doc.add_table(rows=required_rows + 1, cols=n_rev_cols)
        rev_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        rev_table.style = "Table Grid"

        hrow = rev_table.rows[0]
        for ci, h in enumerate(rev_headers):
            cell = hrow.cells[ci]
            _set_cell_bg(cell, _rgb_hex(BLUE_HEADER))
            _set_cell_margins(cell)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(h)
            run.font.bold = True
            run.font.color.rgb = WHITE
            run.font.name = settings.default_font
            run.font.size = Pt(10)

        for ri, revision in enumerate(revision_rows):
            row = rev_table.rows[ri + 1]
            bg = _rgb_hex(LIGHT_BLUE) if ri % 2 == 0 else "FFFFFF"
            row_data = self._revision_row_data(doc_type, ri, revision)
            for ci, val in enumerate(row_data[:n_rev_cols]):
                cell = row.cells[ci]
                _set_cell_bg(cell, bg)
                _set_cell_margins(cell)
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(val)
                run.font.name = settings.default_font
                run.font.size = Pt(10)

        # ── Page break → TOC (page 3, added by add_toc()) ────────────────
        self.doc.add_page_break()

    # ------------------------------------------------------------------
    # Table of Contents
    # ------------------------------------------------------------------

    def add_toc(self):
        toc_heading = self.doc.add_paragraph(style="TOC Heading")
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        toc_heading.add_run("Table of Contents")
        _add_header_border(toc_heading)

        _add_toc_field(self.doc)

        self.doc.add_page_break()

    _CIRCULAR_FONT = "Times New Roman"

    def _circular_run(self, para, text: str, bold: bool = False,
                      underline: bool = False, size: int = 11) -> None:
        run = para.add_run(text)
        run.font.name = self._CIRCULAR_FONT
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.underline = underline
        run.font.color.rgb = BLACK

    def _add_circular_paragraph(self, text: str):
        para = self.doc.add_paragraph()
        _set_paragraph_spacing(para, before=0, after=8)
        self._circular_run(para, text)

    def _add_circular_footer(self):
        """NPCI address in footer — Times New Roman 8 pt, centered."""
        section = self.doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        fpara = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fpara.clear()
        fpara.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fpara.add_run(
            "1001A, The Capital, Bandra Kurla Complex, Bandra (East), Mumbai \u2013 400051"
        )
        fr.font.name = self._CIRCULAR_FONT
        fr.font.size = Pt(8)
        fr.font.color.rgb = DARK_GRAY

    def add_circular_reference_block(self, meta: dict):
        """OC reference number (left) + issue date (right) in a borderless 2-col table."""
        table = self.doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        left = table.rows[0].cells[0]
        right = table.rows[0].cells[1]
        for cell in (left, right):
            _set_cell_margins(cell, top=0, bottom=0, left=0, right=0)

        left_para = left.paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._circular_run(left_para, meta.get("reference_code", ""), bold=True)

        right_para = right.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._circular_run(right_para, meta.get("issue_date", ""), bold=True)

        self.doc.add_paragraph()

    def add_circular_addressee(self, recipient_line: str):
        to_para = self.doc.add_paragraph()
        self._circular_run(to_para, "To,", bold=True)

        recipient_para = self.doc.add_paragraph()
        self._circular_run(recipient_para, recipient_line, bold=True)

    def add_circular_subject(self, subject_line: str):
        para = self.doc.add_paragraph()
        _set_paragraph_spacing(para, before=8, after=10)
        self._circular_run(para, subject_line, bold=True, underline=True)

    def add_circular_signature(self, meta: dict):
        self.doc.add_paragraph()
        close_para = self.doc.add_paragraph()
        close_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._circular_run(close_para, "Yours Sincerely,")

        sd_para = self.doc.add_paragraph()
        self._circular_run(sd_para, "SD/-")

        name_para = self.doc.add_paragraph()
        self._circular_run(
            name_para,
            meta.get("signatory_name", "Authorized Signatory"),
            bold=True,
        )

        for value in (meta.get("signatory_title", ""), meta.get("signatory_department", "")):
            if value:
                p = self.doc.add_paragraph()
                self._circular_run(p, value)

    def add_circular_document(self, plan: dict, sections: list[dict]):
        """Build a Circular with NPCI letterhead, narrow margins, Times New Roman body."""
        # Narrow margins matching Java: top/bottom 0.5", left 0.75", right 0.5"
        sec = self.doc.sections[0]
        sec.top_margin = Inches(0.5)
        sec.bottom_margin = Inches(0.5)
        sec.left_margin = Inches(0.75)
        sec.right_margin = Inches(0.5)

        meta = plan.get("document_meta", {})

        # Letterhead table: org name (left) | empty right cell
        lh_table = self.doc.add_table(rows=1, cols=2)
        lh_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        lh_table.autofit = True
        for cell in lh_table.rows[0].cells:
            _set_cell_margins(cell, top=0, bottom=0, left=0, right=0)
        lh_left = lh_table.rows[0].cells[0]
        lh_para = lh_left.paragraphs[0]
        lh_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._circular_run(
            lh_para,
            meta.get("organization_name", "NATIONAL PAYMENTS CORPORATION OF INDIA"),
            bold=True,
            size=12,
        )

        # Blue ruled separator
        sep = self.doc.add_paragraph()
        _add_header_border(sep)
        _set_paragraph_spacing(sep, before=4, after=4)

        # OC reference + date
        self.add_circular_reference_block(meta)

        # Addressee
        self.add_circular_addressee(meta.get("recipient_line", "All concerned participants"))

        # Subject
        self.add_circular_subject(
            meta.get("subject_line", plan.get("title", "Subject: Circular update"))
        )

        # Body sections
        for sec_item in sections:
            key = sec_item.get("section_key")
            if key in {"letterhead_reference", "addressee_line", "subject_line", "signature_block"}:
                continue
            if key == "dissemination_instruction":
                text = (
                    sec_item.get("paragraphs", [None])[0]
                    or "Please disseminate the information contained herein to the officials concerned."
                )
                self._add_circular_paragraph(text)
                continue

            for para_text in sec_item.get("paragraphs", []):
                if para_text.strip():
                    self._add_circular_paragraph(para_text)
            if sec_item.get("bullet_points"):
                for item in sec_item["bullet_points"]:
                    bp = self.doc.add_paragraph()
                    _set_paragraph_spacing(bp, before=0, after=4)
                    self._circular_run(bp, f"\u2022 {item}")
            if sec_item.get("numbered_items"):
                for i_ni, item in enumerate(sec_item["numbered_items"], 1):
                    np_i = self.doc.add_paragraph()
                    _set_paragraph_spacing(np_i, before=0, after=4)
                    self._circular_run(np_i, f"{i_ni}. {item}")

        # Signature block
        self.add_circular_signature(meta)

        # NPCI address footer
        self._add_circular_footer()

    # ------------------------------------------------------------------
    # Header / Footer
    # ------------------------------------------------------------------

    def add_header_footer(self, doc_title: str, version: str = "1.0", classification: str = "Confidential"):
        section = self.doc.sections[0]

        # Header
        header = section.header
        header.is_linked_to_previous = False
        if header.paragraphs:
            hpara = header.paragraphs[0]
        else:
            hpara = header.add_paragraph()
        hpara.clear()
        hpara.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _add_header_border(hpara)
        run = hpara.add_run(f"{doc_title}  |  v{version}")
        run.font.name = settings.default_font
        run.font.size = Pt(9)
        run.font.color.rgb = BLUE_DARK

        # Explicitly empty the first-page header (cover page gets no header)
        first_page_header = section.first_page_header
        if first_page_header.paragraphs:
            for para in first_page_header.paragraphs:
                para.clear()

        # Footer
        footer = section.footer
        footer.is_linked_to_previous = False
        if footer.paragraphs:
            fpara = footer.paragraphs[0]
        else:
            fpara = footer.add_paragraph()
        fpara.clear()
        fpara.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fpara.add_run(f"{classification}  |  Page ")
        run.font.name = settings.default_font
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = DARK_GRAY
        _add_page_number(fpara)

    # ------------------------------------------------------------------
    # Headings
    # ------------------------------------------------------------------

    def add_heading(self, text: str, level: int = 1):
        para = self.doc.add_paragraph(text, style=f"Heading {min(level, 3)}")
        _set_paragraph_spacing(para, before=18 if level == 1 else 12, after=6)

    # ------------------------------------------------------------------
    # Paragraphs
    # ------------------------------------------------------------------

    def add_paragraph(self, text: str):
        para = self.doc.add_paragraph()
        _set_paragraph_spacing(para, before=0, after=8)
        run = para.add_run(text)
        run.font.name = settings.default_font
        run.font.size = Pt(settings.default_font_size)
        run.font.color.rgb = BLACK

    def add_code_block(self, code: str):
        """Render a code/XML block: Courier New, light gray background, preserved line breaks."""
        para = self.doc.add_paragraph()
        _set_paragraph_spacing(para, before=8, after=8)
        # Set paragraph shading to light gray
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F2F2")
        pPr.append(shd)
        # Add left indent for visual separation
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "360")
        pPr.append(ind)
        lines = code.split("\n")
        for i, line in enumerate(lines):
            run = para.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x1F, 0x2D, 0x3D)
            if i < len(lines) - 1:
                run.add_break()

    # ------------------------------------------------------------------
    # Lists
    # ------------------------------------------------------------------

    def add_bullet_list(self, items: list[str]):
        for item in items:
            para = self.doc.add_paragraph(style="List Bullet")
            _set_paragraph_spacing(para, before=0, after=4)
            run = para.add_run(item)
            run.font.name = settings.default_font
            run.font.size = Pt(settings.default_font_size)

    def add_numbered_list(self, items: list[str]):
        for item in items:
            para = self.doc.add_paragraph(style="List Number")
            _set_paragraph_spacing(para, before=0, after=4)
            run = para.add_run(item)
            run.font.name = settings.default_font
            run.font.size = Pt(settings.default_font_size)

    # ------------------------------------------------------------------
    # Styled table
    # ------------------------------------------------------------------

    def add_styled_table(
        self,
        headers: list[str],
        rows: list[list[str]],
        alignments: list[str] | None = None,
        caption: str = "",
        totals_row: list[str] | None = None,
    ):
        """
        Build a styled table.

        alignments: one string per column — "LEFT" | "CENTER" | "RIGHT".
                    Defaults to LEFT for all columns if not supplied.
        caption:    italic caption paragraph rendered below the table.
        totals_row: optional last row rendered with bold text.
        """
        if not headers:
            return

        _ALIGN_MAP = {
            "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
            "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
            "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
        }

        def _col_align(ci: int) -> WD_ALIGN_PARAGRAPH:
            if alignments and ci < len(alignments):
                return _ALIGN_MAP.get(alignments[ci].upper(), WD_ALIGN_PARAGRAPH.LEFT)
            return WD_ALIGN_PARAGRAPH.LEFT

        n_cols = len(headers)
        all_rows = list(rows)
        if totals_row:
            all_rows.append(totals_row)
        table = self.doc.add_table(rows=len(all_rows) + 1, cols=n_cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        hrow = table.rows[0]
        for ci, h in enumerate(headers):
            cell = hrow.cells[ci]
            _set_cell_bg(cell, _rgb_hex(BLUE_HEADER))
            _set_cell_margins(cell)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(h)
            run.font.bold = True
            run.font.color.rgb = WHITE
            run.font.name = settings.default_font
            run.font.size = Pt(10)

        # Data rows
        for ri, row_data in enumerate(all_rows):
            row = table.rows[ri + 1]
            bg = _rgb_hex(LIGHT_BLUE) if ri % 2 == 0 else "FFFFFF"
            padded = [str(v) for v in row_data[:n_cols]]
            if len(padded) < n_cols:
                padded.extend([""] * (n_cols - len(padded)))
            is_totals = totals_row is not None and ri == len(all_rows) - 1
            for ci, val in enumerate(padded):
                cell = row.cells[ci]
                _set_cell_bg(cell, bg)
                _set_cell_margins(cell)
                para = cell.paragraphs[0]
                para.alignment = _col_align(ci)
                run = para.add_run(str(val))
                run.font.name = settings.default_font
                run.font.size = Pt(10)
                if is_totals:
                    run.font.bold = True

        # Caption
        if caption:
            cap_para = self.doc.add_paragraph()
            _set_paragraph_spacing(cap_para, before=2, after=8)
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = cap_para.add_run(caption)
            cr.italic = True
            cr.font.size = Pt(9)
            cr.font.color.rgb = DARK_GRAY
        else:
            self.doc.add_paragraph()

    # ------------------------------------------------------------------
    # Diagram / image
    # ------------------------------------------------------------------

    def add_diagram(self, image_path: str, caption: str = ""):
        try:
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_paragraph_spacing(para, before=12, after=6)
            run = para.add_run()
            run.add_picture(image_path, width=Inches(6.5))
        except Exception as e:
            logger.warning("Could not embed image %s: %s", image_path, e)
            para = self.doc.add_paragraph()
            run = para.add_run(f"[Diagram: {caption or image_path}]")
            run.font.italic = True
            run.font.color.rgb = DARK_GRAY

        if caption:
            cap_para = self.doc.add_paragraph()
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_paragraph_spacing(cap_para, before=0, after=12)
            run = cap_para.add_run(f"Figure: {caption}")
            run.font.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = DARK_GRAY

    def add_diagram_at_heading(self, image_path: str, target_heading: str, caption: str = ""):
        """
        Scan the document for a heading whose text matches target_heading
        (case-insensitive) and insert the image + caption immediately after it.
        Falls back to appending at the end if heading is not found.
        """
        paragraphs = self.doc.paragraphs
        idx = next(
            (i for i, p in enumerate(paragraphs)
             if p.text.strip().lower() == target_heading.strip().lower()),
            None,
        )

        # Build image paragraph (appended at end first, then moved)
        img_para = self.doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(img_para, before=12, after=6)
        try:
            img_para.add_run().add_picture(image_path, width=Inches(6.5))
        except Exception as e:
            logger.warning("Could not embed image %s: %s", image_path, e)
            img_para.clear()
            r = img_para.add_run(f"[Diagram: {caption or image_path}]")
            r.font.italic = True
            r.font.color.rgb = DARK_GRAY

        # Build caption paragraph
        cap_para = self.doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(cap_para, before=0, after=12)
        cr = cap_para.add_run(f"Figure: {caption}" if caption else "")
        cr.italic = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = DARK_GRAY

        # Move both paragraphs to right after the target heading
        if idx is not None:
            _move_paragraph_after(self.doc, img_para, idx)
            _move_paragraph_after(self.doc, cap_para, idx + 1)

    # ------------------------------------------------------------------
    # TSD API spec section
    # ------------------------------------------------------------------

    def add_api_spec_section(
        self,
        api_label: str,
        purpose: str,
        xml_samples: list[dict],
        has_tag_rows: bool = False,
    ):
        """
        Render one TSD API spec block:
          bold+underlined API label → purpose bullets → R&R heading →
          (italic+blue label + Courier XML per sample) → New/Modified Tags heading.
        The caller is responsible for injecting the R&R and tag tables after the headings.
        """
        # API label — bold + underlined, Calibri 12 pt
        p = self.doc.add_paragraph()
        _set_paragraph_spacing(p, before=14, after=6)
        run = p.add_run(api_label)
        run.bold = True
        run.underline = True
        run.font.name = "Calibri"
        run.font.size = Pt(12)
        self.doc.add_paragraph()

        # Purpose bullets
        for line in purpose.split("\n"):
            line = line.strip().lstrip("•").strip()
            if line:
                bp = self.doc.add_paragraph()
                _set_paragraph_spacing(bp, before=0, after=4)
                bp.paragraph_format.left_indent = Inches(0.5)
                r = bp.add_run(f"\u2022 {line}")
                r.font.name = "Calibri"
                r.font.size = Pt(11)
        self.doc.add_paragraph()

        # Roles and responsibilities heading
        rr = self.doc.add_paragraph()
        _set_paragraph_spacing(rr, before=6, after=4)
        rr_run = rr.add_run("Roles and responsibilities:")
        rr_run.bold = True
        rr_run.font.name = "Calibri"
        rr_run.font.size = Pt(11)
        self.doc.add_paragraph()

        # XML / request–response samples
        for sample in xml_samples:
            label_val = sample.get("label", "")
            if label_val:
                label_p = self.doc.add_paragraph()
                _set_paragraph_spacing(label_p, before=6, after=4)
                lr = label_p.add_run(label_val)
                lr.italic = True
                lr.bold = True
                lr.font.color.rgb = RGBColor(0x1F, 0x3B, 0x64)
                lr.font.name = "Calibri"
                lr.font.size = Pt(11)

            xml_content = sample.get("xml", "")
            if xml_content:
                xml_p = self.doc.add_paragraph()
                _set_paragraph_spacing(xml_p, before=0, after=6)
                xml_p.paragraph_format.left_indent = Inches(0.25)
                # Set gray shading
                pPr = xml_p._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "F2F2F2")
                pPr.append(shd)
                xml_run = xml_p.add_run(xml_content)
                xml_run.font.name = "Courier New"
                xml_run.font.size = Pt(9)
            self.doc.add_paragraph()

        # New / Modified Tags heading (table appended by caller)
        if has_tag_rows:
            tag_h = self.doc.add_paragraph()
            _set_paragraph_spacing(tag_h, before=6, after=4)
            tag_run = tag_h.add_run("New / Modified Tags:")
            tag_run.bold = True
            tag_run.font.name = "Calibri"
            tag_run.font.size = Pt(11)
            self.doc.add_paragraph()

    # ------------------------------------------------------------------
    # Error codes injection
    # ------------------------------------------------------------------

    def add_error_codes_at_heading(self, error_rows: list, notes: list | None = None):
        """
        Append the error-codes table and optional notes.
        If an 'Error Handling' or 'Error Code' heading exists in the document
        it will be found and the table rendered just after (via section content
        rendering — positional move is left to caller ordering).
        """
        headers = ["Response Code", "Error Code", "Description", "API", "Entity", "TD/BD"]
        alignments = ["CENTER", "CENTER", "LEFT", "LEFT", "LEFT", "CENTER"]
        self.add_styled_table(headers, error_rows, alignments=alignments)

        if notes:
            notes_h = self.doc.add_paragraph()
            _set_paragraph_spacing(notes_h, before=8, after=4)
            notes_h.add_run("Notes:").bold = True
            for i_n, note in enumerate(notes):
                np = self.doc.add_paragraph()
                _set_paragraph_spacing(np, before=0, after=4)
                np.paragraph_format.left_indent = Inches(0.25)
                np.add_run(f"{i_n + 1}. {note}")

    # ------------------------------------------------------------------
    # Section content renderer
    # ------------------------------------------------------------------

    def add_section_content(self, content: GeneratedContent):
        self.add_heading(content.section_heading, level=content.level)

        for para_text in content.paragraphs:
            if para_text.strip():
                for chunk in _split_long_paragraph(para_text):
                    self.add_paragraph(chunk)

        if content.bullet_points:
            self.add_bullet_list(content.bullet_points)

        if content.numbered_items:
            self.add_numbered_list(content.numbered_items)

        if content.table_data and content.table_data.headers:
            self.add_styled_table(content.table_data.headers, content.table_data.rows)

        for code in getattr(content, "code_blocks", []):
            if code.strip():
                self.add_code_block(code)

        if content.diagram_path and Path(content.diagram_path).exists():
            self.add_diagram(content.diagram_path, caption=content.section_heading)

    # ------------------------------------------------------------------
    # Save
    # ------------------------------------------------------------------

    def save(self, output_path: str) -> str:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(output_path)
        logger.info("Document saved: %s", output_path)
        return output_path


# ---------------------------------------------------------------------------
# High-level assembler function
# ---------------------------------------------------------------------------

def assemble_document(
    plan: dict,
    sections: list[dict],
    output_path: str,
    version: str = "1.0",
    diagram_specs: list[dict] | None = None,
    generated_diagrams: dict[str, str] | None = None,
) -> str:
    """
    Build the final .docx from a document plan and generated sections.

    Phase 1 — create base document: cover + revision history + TOC + section content.
    Phase 2 — TSD only: inject API spec blocks (label + XML + R&R + tags).
    Phase 3 — embed diagrams at their target headings using add_diagram_at_heading().

    Returns the output file path.
    """
    builder = DocxBuilder()

    title = plan.get("title", "Document")
    subtitle = plan.get("subtitle", "")
    doc_type = plan.get("doc_type", "BRD")
    meta = plan.get("document_meta", {})
    include_cover_page = plan.get("include_cover_page", True)
    include_toc = plan.get("include_toc", True)

    # ── Circular: entirely different rendering path ───────────────────────
    if doc_type.lower() == "circular":
        builder.add_circular_document(plan, sections)
        return builder.save(output_path)

    # ── Phase 1: base document ────────────────────────────────────────────
    if include_cover_page:
        builder.add_cover_page(title, subtitle, doc_type, version, meta=meta)

    classification = meta.get("classification", "Confidential")
    builder.add_header_footer(title, version, classification=classification)

    if include_toc:
        builder.add_toc()

    for sec_dict in sections:
        content = GeneratedContent(
            section_key=sec_dict.get("section_key"),
            section_heading=sec_dict.get("section_heading", "Section"),
            render_style=sec_dict.get("render_style", "body"),
            paragraphs=sec_dict.get("paragraphs", []),
            bullet_points=sec_dict.get("bullet_points", []),
            numbered_items=sec_dict.get("numbered_items", []),
            table_data=TableData(**sec_dict["table_data"]) if sec_dict.get("table_data") else None,
            diagram_path=None,   # diagrams injected in Phase 3; skip inline
            level=sec_dict.get("level", 1),
            code_blocks=sec_dict.get("code_blocks", []),
        )
        builder.add_section_content(content)

    # ── Phase 2 (TSD only): API spec blocks ──────────────────────────────
    if doc_type.strip().lower() == "tsd":
        tsd_meta = meta.get("tsdMetadata") or {}
        api_specs = tsd_meta.get("apiSpecs") or []
        for api in api_specs:
            rr_rows = api.get("rrRows") or []
            tag_rows = api.get("tagRows") or []
            xml_samples = api.get("xmlSamples") or []
            builder.add_api_spec_section(
                api_label=api.get("apiLabel", api.get("apiName", "API")),
                purpose=api.get("purpose", ""),
                xml_samples=xml_samples,
                has_tag_rows=bool(tag_rows),
            )
            if rr_rows:
                builder.add_styled_table(
                    ["Step", "Activity", "Responsible"],
                    rr_rows,
                    alignments=["LEFT", "LEFT", "CENTER"],
                    caption="Roles and Responsibilities",
                )
            if tag_rows:
                builder.add_styled_table(
                    ["Tag Num", "Message Items", "<xml tag>", "Occurrence"],
                    tag_rows,
                    alignments=["CENTER", "LEFT", "LEFT", "CENTER"],
                    caption=f"New / Modified Tags — {api.get('apiName', '')}",
                )

        # Error codes
        error_rows = tsd_meta.get("errorRows") or []
        notes = tsd_meta.get("notes") or []
        if error_rows:
            builder.add_error_codes_at_heading(error_rows, notes)

    # ── Phase 3: embed diagrams at target headings ────────────────────────
    if diagram_specs and generated_diagrams:
        for spec in diagram_specs:
            did = spec.get("diagram_id", "")
            png_path = generated_diagrams.get(did, "")
            if not png_path or not Path(png_path).exists():
                continue
            target_heading = spec.get("target_heading", "")
            caption = spec.get("description", spec.get("caption", ""))
            if target_heading:
                builder.add_diagram_at_heading(png_path, target_heading, caption)
            else:
                # Fallback: append at end
                builder.add_diagram(png_path, caption)

    return builder.save(output_path)
