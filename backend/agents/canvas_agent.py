"""
Agent 3 — Product Canvas Generator (LangGraph)

Generates a 10-section NPCI Product Canvas matching the V1.2 Build Framework template:
  1. Feature
  2. Need
  3. Market View
  4. Scalability
  5. Validation
  6. Product Operating
  7. Product Comms
  8. Pricing
  9. Potential Risks
  10. Compliance
"""
import json
import logging
import io
from typing import TypedDict, Annotated, Optional
import operator
from pathlib import Path

from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage, SystemMessage
from langgraph.graph import StateGraph, END
from langgraph.checkpoint.memory import MemorySaver

from agents.rag_client import query_rag_sync
import config

logger = logging.getLogger(__name__)

# --------------------------------------------------------------------------
# Template section definitions — exactly matching V1.2_Canvas_ProductBuild
# --------------------------------------------------------------------------

CANVAS_SECTIONS = [
    {
        "key": "feature",
        "title": "1. Feature",
        "subtitle": "Explain the feature for a layman",
    },
    {
        "key": "need",
        "title": "2. Need",
        "subtitle": "Why should we do this?",
        "subsections": [
            "Why should we do this?",
            "Differentiation (incremental or exponential)",
            "Delta in user experience",
            "What will it cannibalize?",
            "What if we don't build this?",
        ],
    },
    {
        "key": "market_view",
        "title": "3. Market View",
        "subsections": [
            "Ecosystem anticipated (informal) response",
            "Ecosystem efforts (costs to make this work)",
            "Anticipated regulatory view",
        ],
    },
    {
        "key": "scalability",
        "title": "4. Scalability",
        "subsections": [
            "Market anchors to make it big? (demand and supply)",
            "Impact opportunity (users, delta in time, revenue)",
        ],
    },
    {
        "key": "validation",
        "title": "5. Validation",
        "subsections": [
            "Creating and operating MVP",
            "Data it will generate to create insights",
        ],
    },
    {
        "key": "product_operating",
        "title": "6. Product Operating",
        "subsections": [
            "3 Success KPIs",
            "Grievance redressal (Trust)",
            "Day 0 automation",
            "Impact on SGF",
            "Impact on FRM",
            "Impact on existing txns and infra",
        ],
    },
    {
        "key": "product_comms",
        "title": "7. Product Comms (external + internal)",
        "subsections": [
            "Product demo (polished version of MVP)",
            "Product video",
            "Explanation video by PM",
            "FAQs + trained LLM",
            "Circular",
            "Product doc (w/ Specs, explanations, test cases/UI/UX guidelines)",
        ],
    },
    {
        "key": "pricing",
        "title": "8. Pricing",
        "subsections": [
            "3 Year view of pricing & revenue",
            "Market ability to pay the price (total pie)",
            "Market view to pay the price",
        ],
    },
    {
        "key": "potential_risks",
        "title": "9. Potential Risks",
        "subsections": [
            "Fraud risk",
            "Infosec risk",
            "Legal risk",
            "Data privacy risk",
            "2nd order negative effect",
        ],
    },
    {
        "key": "compliance",
        "title": "10. Compliance",
        "subsections": [
            "Existing guideline change",
            "New guideline addition",
            "Must have compliances in NPCI product circular for ecosystem",
        ],
    },
]

CANVAS_SECTION_KEYS = [s["key"] for s in CANVAS_SECTIONS]
CANVAS_TITLE_MAP = {s["key"]: s["title"] for s in CANVAS_SECTIONS}

SECTION_PROMPTS = {
    "feature": """Write a clear, jargon-free explanation of this feature for a layman (non-technical audience).
Describe what the feature does, how a user would experience it, and why it matters in 2–3 short paragraphs.
Avoid technical terms. Write as if explaining to a senior executive unfamiliar with the domain.""",

    "need": """Answer ALL of the following for the product canvas Need section. Use the exact sub-headings:

**Why should we do this?**
[Justify the business case — market gap, user pain point, or strategic opportunity]

**Differentiation (incremental or exponential)**
[State whether this is an incremental improvement or an exponential leap. Explain why.]

**Delta in user experience**
[Describe the concrete improvement in user experience compared to today]

**What will it cannibalize?**
[Identify existing products/volumes/revenues this feature may cannibalize]

**What if we don't build this?**
[State the cost of inaction — competitive risk, user loss, market position]""",

    "market_view": """Answer ALL of the following for the product canvas Market View section. Use the exact sub-headings:

**Ecosystem anticipated (informal) response**
[How will banks, PSPs, merchants, and other ecosystem participants likely react?]

**Ecosystem efforts (costs to make this work)**
[What effort/investment does the ecosystem need to make to adopt this?]

**Anticipated regulatory view**
[How is RBI/regulatory body likely to view this? Any approvals needed?]""",

    "scalability": """Answer ALL of the following for the product canvas Scalability section. Use the exact sub-headings:

**Market anchors to make it big? (demand and supply)**
[What demand-side and supply-side anchors will drive scale? Who are the key players?]

**Impact opportunity (users, delta in time, revenue)**
[Quantify the opportunity: projected user base, time savings, revenue potential over 3 years]""",

    "validation": """Answer ALL of the following for the product canvas Validation section. Use the exact sub-headings:

**Creating and operating MVP**
[Describe the minimum viable product — what is included, who builds it, how it is operated in pilot]

**Data it will generate to create insights**
[What data will the MVP generate? How will this data inform the go/no-go decision for full launch?]""",

    "product_operating": """Answer ALL of the following for the product canvas Product Operating section. Use the exact sub-headings:

**3 Success KPIs**
[List exactly 3 measurable KPIs with targets and timelines]

**Grievance redressal (Trust)**
[How will user complaints/disputes be handled? What is the SLA?]

**Day 0 automation**
[What processes are automated from day one? What is manual initially?]

**Impact on SGF**
[How does this feature impact the Settlement and Guarantee Fund?]

**Impact on FRM**
[How does this feature impact Fraud and Risk Management systems?]

**Impact on existing txns and infra**
[What is the impact on current transaction flows, volumes, and technical infrastructure?]""",

    "product_comms": """Answer ALL of the following for the product canvas Product Comms section. Use the exact sub-headings:

**Product demo (polished version of MVP)**
[Describe the demo format, target audience, and key scenarios to showcase]

**Product video**
[Describe the product explainer video — format, duration, target audience]

**Explanation video by PM**
[Describe the PM walkthrough video — key messages, format, internal vs external]

**FAQs + trained LLM**
[List top 5 FAQs. Describe how the LLM will be trained to handle queries.]

**Circular**
[Describe the NPCI circular — scope, mandatory requirements for ecosystem, timeline]

**Product doc (w/ Specs, explanations, test cases/UI/UX guidelines)**
[Outline the structure of the product documentation — key sections, who maintains it]""",

    "pricing": """Answer ALL of the following for the product canvas Pricing section. Use the exact sub-headings:

**3 Year view of pricing & revenue**
[Provide a structured Year 1 / Year 2 / Year 3 revenue projection with pricing model]

**Market ability to pay the price (total pie)**
[What is the total addressable market in terms of revenue? What is the ecosystem's overall capacity to pay?]

**Market view to pay the price**
[How do participants view the pricing — is it fair, too high, too low? Any benchmarks?]""",

    "potential_risks": """Answer ALL of the following for the product canvas Potential Risks section. Use the exact sub-headings:

**Fraud risk**
[What fraud vectors does this feature introduce? Mitigation strategies?]

**Infosec risk**
[What information security risks exist? How are they mitigated?]

**Legal risk**
[What legal/contractual risks exist? How are they managed?]

**Data privacy risk**
[What personal data is processed? What are the privacy risks and controls?]

**2nd order negative effect**
[What unintended negative consequences could emerge at scale?]""",

    "compliance": """Answer ALL of the following for the product canvas Compliance section. Use the exact sub-headings:

**Existing guideline change**
[Which existing RBI/NPCI guidelines need to be amended? What changes are required?]

**New guideline addition**
[What new guidelines or circulars need to be issued? What should they cover?]

**Must have compliances in NPCI product circular for ecosystem**
[List the mandatory compliance requirements that all ecosystem participants must meet as specified in the NPCI circular]""",
}


# --------------------------------------------------------------------------
# LangGraph state
# --------------------------------------------------------------------------

class CanvasState(TypedDict):
    research_report: dict
    requirement_output: dict
    sections: Annotated[dict, lambda a, b: {**a, **b}]
    section_versions: dict
    canvas_versions: list[dict]
    current_version: int
    target_section: Optional[str]
    regen_instructions: Optional[str]
    status: str


# --------------------------------------------------------------------------
# LLM
# --------------------------------------------------------------------------

def _llm():
    return ChatOpenAI(
        model=config.MODEL_NAME,
        api_key=config.OPENAI_API_KEY,
        temperature=0.3,
    )


# --------------------------------------------------------------------------
# Core generation
# --------------------------------------------------------------------------

def _generate_single_section(
    key: str,
    research_report: dict,
    requirement_output: dict,
    existing_content: Optional[str] = None,
    instructions: Optional[str] = None,
) -> str:
    """Generate one canvas section (sync)."""
    llm = _llm()
    section_meta = next((s for s in CANVAS_SECTIONS if s["key"] == key), {})
    title = section_meta.get("title", key)
    prompt = SECTION_PROMPTS.get(key, f"Write the {title} section of the product canvas.")

    research_context = research_report.get("summary", "")
    for sec in research_report.get("sections", []):
        research_context += f"\n\n### {sec.get('title', '')}\n{sec.get('content', '')}"

    feature_request = requirement_output.get("feature_request", "")
    req_json = json.dumps(requirement_output, indent=2)[:1500]

    parts = [
        f"You are writing the **{title}** section of an NPCI Product Canvas (Build Framework V1.2).",
        f"\nFeature: {feature_request}",
        f"\n\n--- Research Context ---\n{research_context[:4000]}",
        f"\n\n--- Structured Requirements ---\n{req_json}",
        f"\n\n--- Your Task ---\n{prompt}",
    ]

    if existing_content and instructions:
        parts.append(f"\n\n--- Existing Content (refine this) ---\n{existing_content}")
        parts.append(f"\n\n--- Refinement Instructions ---\n{instructions}")

    parts.append("\n\nBe specific, data-driven, and concise. Use the NPCI/UPI payments context. Format with the sub-headings exactly as specified.")

    system = SystemMessage(content=(
        "You are a senior product manager at NPCI writing a Product Canvas document. "
        "You are expert in Indian payments, UPI, RBI regulations, and NPCI ecosystem. "
        "Be concise, structured, and specific. Use markdown for sub-headings and bullet points."
    ))
    human = HumanMessage(content="\n".join(parts))

    response = llm.invoke([system, human])
    return response.content.strip()


# --------------------------------------------------------------------------
# Graph nodes
# --------------------------------------------------------------------------

def generate_canvas(state: CanvasState) -> dict:
    """Generate all 10 canvas sections (sync)."""
    research_report = state["research_report"]
    requirement_output = state["requirement_output"]

    new_sections = {}
    new_section_versions = {}

    for sec in CANVAS_SECTIONS:
        key = sec["key"]
        content = _generate_single_section(key, research_report, requirement_output)
        new_sections[key] = content
        new_section_versions[key] = 1
        logger.info(f"Generated canvas section: {key}")

    new_version = 1
    canvas_snapshot = _build_snapshot(new_sections, new_section_versions, new_version)

    return {
        "sections": new_sections,
        "section_versions": new_section_versions,
        "canvas_versions": [canvas_snapshot],
        "current_version": new_version,
        "status": "ready",
        "target_section": None,
        "regen_instructions": None,
    }


def regenerate_section(state: CanvasState) -> dict:
    """Regenerate a single section (sync)."""
    key = state["target_section"]
    if not key or key not in CANVAS_SECTION_KEYS:
        logger.warning(f"Invalid section key for regeneration: {key}")
        return {"status": "ready"}

    content = _generate_single_section(
        key=key,
        research_report=state["research_report"],
        requirement_output=state["requirement_output"],
        existing_content=state["sections"].get(key),
        instructions=state.get("regen_instructions"),
    )

    new_sections = {**state["sections"], key: content}
    new_sv = {**state.get("section_versions", {}), key: state.get("section_versions", {}).get(key, 1) + 1}
    new_version = state["current_version"] + 1
    canvas_snapshot = _build_snapshot(new_sections, new_sv, new_version)
    versions = list(state.get("canvas_versions", []))
    versions.append(canvas_snapshot)

    return {
        "sections": new_sections,
        "section_versions": new_sv,
        "canvas_versions": versions,
        "current_version": new_version,
        "status": "ready",
        "target_section": None,
        "regen_instructions": None,
    }


def update_section_manual(state: CanvasState, key: str, content: str) -> dict:
    new_sections = {**state["sections"], key: content}
    new_sv = {**state.get("section_versions", {}), key: state.get("section_versions", {}).get(key, 1) + 1}
    new_version = state["current_version"] + 1
    canvas_snapshot = _build_snapshot(new_sections, new_sv, new_version)
    versions = list(state.get("canvas_versions", []))
    versions.append(canvas_snapshot)
    return {
        **state,
        "sections": new_sections,
        "section_versions": new_sv,
        "canvas_versions": versions,
        "current_version": new_version,
    }


def _build_snapshot(sections: dict, section_versions: dict, version: int) -> dict:
    return {
        "version": version,
        "sections": [
            {
                "key": k,
                "title": CANVAS_TITLE_MAP[k],
                "content": sections.get(k, ""),
                "version": section_versions.get(k, 1),
            }
            for k in CANVAS_SECTION_KEYS
        ],
    }


# --------------------------------------------------------------------------
# Graph
# --------------------------------------------------------------------------

def build_canvas_graph() -> StateGraph:
    workflow = StateGraph(CanvasState)
    workflow.add_node("generate_canvas", generate_canvas)
    workflow.add_node("regenerate_section", regenerate_section)
    workflow.set_entry_point("generate_canvas")
    workflow.add_edge("generate_canvas", END)
    workflow.add_edge("regenerate_section", END)
    return workflow


_checkpointer = None
_canvas_graph = None

def _get_graph():
    global _checkpointer, _canvas_graph
    if _canvas_graph is None:
        _checkpointer = MemorySaver()
        _canvas_graph = build_canvas_graph().compile(checkpointer=_checkpointer)
    return _canvas_graph


# --------------------------------------------------------------------------
# Public API
# --------------------------------------------------------------------------

def generate_product_canvas(
    session_id: str,
    research_report: dict,
    requirement_output: dict,
) -> dict:
    """Generate full product canvas (sync, called via asyncio.to_thread)."""
    config_dict = {"configurable": {"thread_id": f"canvas_{session_id}"}}
    initial_state = {
        "research_report": research_report,
        "requirement_output": requirement_output,
        "sections": {},
        "section_versions": {},
        "canvas_versions": [],
        "current_version": 0,
        "target_section": None,
        "regen_instructions": None,
        "status": "generating",
    }
    return _get_graph().invoke(initial_state, config=config_dict)


def regenerate_canvas_section(
    session_id: str,
    section_key: str,
    instructions: Optional[str] = None,
) -> dict:
    """Regenerate one canvas section (sync, called via asyncio.to_thread)."""
    config_dict = {"configurable": {"thread_id": f"canvas_{session_id}"}}
    snapshot = _get_graph().get_state(config_dict)
    current = snapshot.values
    updated = {
        **current,
        "target_section": section_key,
        "regen_instructions": instructions,
    }
    return _get_graph().invoke(updated, config=config_dict)


def update_canvas_section_manually(
    session_id: str,
    section_key: str,
    content: str,
) -> dict:
    """Manually update one canvas section (sync)."""
    config_dict = {"configurable": {"thread_id": f"canvas_{session_id}"}}
    snapshot = _get_graph().get_state(config_dict)
    current = dict(snapshot.values)
    updated = update_section_manual(current, section_key, content)
    _get_graph().update_state(config_dict, updated)
    return updated


# --------------------------------------------------------------------------
# Export helpers — V1.2 template layout
# --------------------------------------------------------------------------

# Grid layout matching the template:
# Row 1: feature (full width)
# Row 2: need | market_view | scalability
# Row 3: validation | product_operating | product_comms
# Row 4: pricing | potential_risks | compliance

_GRID_LAYOUT = [
    ["feature"],                                          # row 1 (full width)
    ["need", "market_view", "scalability"],               # row 2
    ["validation", "product_operating", "product_comms"], # row 3
    ["pricing", "potential_risks", "compliance"],          # row 4
]


def export_docx(sections: dict, feature_title: str = "") -> bytes:
    """Generate DOCX matching the V1.2 template grid layout."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import re

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # Header: "Build framework for ___"
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header_para.add_run(f"Build framework for  {feature_title or '_______________'}")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)

    def _set_cell_bg(cell, hex_color="FFFFFF"):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def _write_cell(cell, key):
        meta = next((s for s in CANVAS_SECTIONS if s["key"] == key), {})
        title = meta.get("title", key)
        content = sections.get(key, "")

        # Section title
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)

        # Subtitle if present
        subtitle = meta.get("subtitle")
        if subtitle:
            sp = cell.add_paragraph()
            sr = sp.add_run(subtitle)
            sr.italic = True
            sr.font.size = Pt(8)
            sr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        cell.add_paragraph()  # spacer

        # Content — parse markdown lightly
        for line in content.split("\n"):
            line_stripped = line.strip()
            if not line_stripped:
                cell.add_paragraph()
                continue
            if line_stripped.startswith("**") and line_stripped.endswith("**"):
                cp = cell.add_paragraph()
                cr = cp.add_run(line_stripped[2:-2])
                cr.bold = True
                cr.font.size = Pt(8)
                cr.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            elif line_stripped.startswith("- ") or line_stripped.startswith("* "):
                cp = cell.add_paragraph(style="List Bullet")
                text = re.sub(r"\*\*(.+?)\*\*", r"\1", line_stripped[2:])
                cr = cp.add_run(text)
                cr.font.size = Pt(8)
            else:
                cp = cell.add_paragraph()
                parts_bold = re.split(r"(\*\*.+?\*\*)", line_stripped)
                for part in parts_bold:
                    if part.startswith("**") and part.endswith("**"):
                        cr = cp.add_run(part[2:-2])
                        cr.bold = True
                        cr.font.size = Pt(8)
                    else:
                        cr = cp.add_run(part)
                        cr.font.size = Pt(8)

    # Row 1: feature (full-width, single-column table)
    t1 = doc.add_table(rows=1, cols=1)
    t1.style = "Table Grid"
    cell = t1.rows[0].cells[0]
    _write_cell(cell, "feature")
    _set_cell_bg(cell, "F8FAFF")
    doc.add_paragraph()

    # Rows 2–4: three-column tables
    for row_keys in _GRID_LAYOUT[1:]:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        for i, key in enumerate(row_keys):
            cell = table.rows[0].cells[i]
            _write_cell(cell, key)
            if i == 0:
                _set_cell_bg(cell, "FFFFFF")
            elif i == 1:
                _set_cell_bg(cell, "FAFAFA")
            else:
                _set_cell_bg(cell, "F0F4FA")
        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def export_pdf(sections: dict, feature_title: str = "") -> bytes:
    """Generate PDF matching the V1.2 template grid layout using ReportLab (pure-Python)."""
    import io
    import re as _re
    from reportlab.lib.pagesizes import A3, landscape
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    )
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT

    NPCI_BLUE = colors.HexColor("#1B3F8F")
    BORDER = colors.HexColor("#cccccc")
    BG_LIGHT = colors.HexColor("#f4f7fb")
    BG_WHITE = colors.white

    title_style = ParagraphStyle(
        "CellTitle", fontName="Helvetica-Bold", fontSize=7,
        textColor=NPCI_BLUE, leading=9, spaceAfter=2,
    )
    subtitle_style = ParagraphStyle(
        "CellSub", fontName="Helvetica-Oblique", fontSize=6,
        textColor=colors.HexColor("#888888"), leading=8, spaceAfter=2,
    )
    body_style = ParagraphStyle(
        "CellBody", fontName="Helvetica", fontSize=6.5,
        textColor=colors.HexColor("#333333"), leading=9,
        spaceAfter=1,
    )
    bold_body_style = ParagraphStyle(
        "CellBold", fontName="Helvetica-Bold", fontSize=6.5,
        textColor=colors.HexColor("#1a3c6e"), leading=9, spaceAfter=1,
    )

    def _md_to_paras(content: str) -> list:
        """Lightweight markdown → list of Paragraph objects."""
        paras = []
        for line in content.split("\n"):
            s = line.strip()
            if not s:
                paras.append(Spacer(1, 3))
                continue
            # heading — treat as bold body
            if s.startswith("#"):
                s = _re.sub(r"^#+\s*", "", s)
                s = _re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", s)
                paras.append(Paragraph(s, bold_body_style))
            # bullet
            elif s.startswith("- ") or s.startswith("* "):
                text = _re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", s[2:])
                paras.append(Paragraph(f"• {text}", body_style))
            # bold-only line
            elif s.startswith("**") and s.endswith("**"):
                paras.append(Paragraph(s[2:-2], bold_body_style))
            else:
                text = _re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", s)
                paras.append(Paragraph(text, body_style))
        return paras

    def _cell_content(key: str) -> list:
        meta = next((sec for sec in CANVAS_SECTIONS if sec["key"] == key), {})
        title = meta.get("title", key)
        subtitle = meta.get("subtitle", "")
        content = sections.get(key, "")
        items = [Paragraph(title, title_style)]
        if subtitle:
            items.append(Paragraph(subtitle, subtitle_style))
        items.append(Spacer(1, 3))
        items.extend(_md_to_paras(content))
        return items

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=landscape(A3),
        leftMargin=15 * mm, rightMargin=15 * mm,
        topMargin=15 * mm, bottomMargin=15 * mm,
    )

    page_w = landscape(A3)[0] - 30 * mm  # usable width

    cell_style_base = [
        ("BOX", (0, 0), (-1, -1), 0.5, BORDER),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]

    story = []

    # Page header
    hdr_style = ParagraphStyle(
        "Hdr", fontName="Helvetica-Bold", fontSize=12,
        textColor=NPCI_BLUE, spaceAfter=6,
    )
    story.append(Paragraph(
        f"Build framework for &nbsp; {feature_title or '_______________'}",
        hdr_style,
    ))

    # Row 1 — feature (full width)
    row1 = Table(
        [[_cell_content("feature")]],
        colWidths=[page_w],
        repeatRows=0,
    )
    row1.setStyle(TableStyle(cell_style_base + [
        ("BACKGROUND", (0, 0), (0, 0), BG_LIGHT),
    ]))
    story.append(row1)
    story.append(Spacer(1, 3))

    # Rows 2–4 — three columns each
    col_w = page_w / 3
    bg_cycle = [BG_WHITE, colors.HexColor("#fafafa"), colors.HexColor("#f0f4fa")]
    for row_keys in _GRID_LAYOUT[1:]:
        data = [[_cell_content(k) for k in row_keys]]
        tbl = Table(data, colWidths=[col_w] * 3)
        row_style = cell_style_base + [
            ("INNERGRID", (0, 0), (-1, -1), 0.5, BORDER),
        ]
        for i, bg in enumerate(bg_cycle):
            row_style.append(("BACKGROUND", (i, 0), (i, 0), bg))
        tbl.setStyle(TableStyle(row_style))
        story.append(tbl)
        story.append(Spacer(1, 3))

    doc.build(story)
    buf.seek(0)
    return buf.read()
